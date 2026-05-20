from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\DELL\Downloads\i-fall")
OUT = ROOT / "rapports" / "guide-explication-presentation-TP-0.docx"

SKILL_SCRIPTS = Path(
    r"C:\Users\DELL\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
sys.path.append(str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights

CONTENT_WIDTH_DXA = 10080

NAVY = "0B1F3A"
TEAL = "0F7182"
BLUE = "2563EB"
GREEN = "0F766E"
ORANGE = "EA580C"
RED = "B91C1C"
GRAY = "64748B"
TEXT = "111827"
BORDER = "CBD5E1"
LIGHT_BLUE = "EFF6FF"
LIGHT_TEAL = "E6FFFA"
LIGHT_GREEN = "ECFDF5"
LIGHT_ORANGE = "FFF7ED"
LIGHT_RED = "FEF2F2"
LIGHT_GRAY = "F8FAFC"


def set_font(run, size=10.2, bold=False, italic=False, color=TEXT, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def pspace(paragraph, before=0, after=3, line=1.04):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcBorders")
    if node is None:
        node = OxmlElement("w:tcBorders")
        tc_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = node.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            node.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), size)
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def cell_margins(cell, top=95, start=130, bottom=95, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = tc_mar.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def add_field(paragraph, instr):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(text)
    run._r.append(separate)
    run._r.append(fallback)
    run._r.append(end)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.35)
    sec.bottom_margin = Cm(1.35)
    sec.left_margin = Cm(1.55)
    sec.right_margin = Cm(1.55)
    sec.header_distance = Cm(0.55)
    sec.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size, color in [
        ("Title", 22, NAVY),
        ("Subtitle", 11, GRAY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11.2, BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("Guide oral - Projet ChatUser / TP 0")
    set_font(r, size=8.5, bold=True, color=GRAY)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Page ")
    set_font(r, size=8.4, color=GRAY)
    add_field(footer, "PAGE")
    r = footer.add_run(" / ")
    set_font(r, size=8.4, color=GRAY)
    add_field(footer, "NUMPAGES")


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    pspace(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    pspace(p, before=6, after=3)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        pspace(p, after=2)
        r = p.add_run(item)
        set_font(r, size=9.9)


def add_callout(doc, title, body, kind="info"):
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "oral": (LIGHT_TEAL, TEAL),
        "warn": (LIGHT_ORANGE, ORANGE),
        "ok": (LIGHT_GREEN, GREEN),
        "danger": (LIGHT_RED, RED),
    }
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell, accent, "6")
    cell_margins(cell, top=135, start=190, bottom=130, end=190)
    p = cell.paragraphs[0]
    pspace(p, after=2)
    r = p.add_run(title)
    set_font(r, size=10.0, bold=True, color=accent)
    p = cell.add_paragraph()
    pspace(p, after=0)
    r = p.add_run(body)
    set_font(r, size=9.5)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def add_table(doc, headers, rows, weights=None, header_fill=LIGHT_GRAY, font_size=8.6):
    if weights is None:
        weights = [1] * len(headers)
    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, header_fill)
        borders(cell)
        cell_margins(cell)
        for p in cell.paragraphs:
            pspace(p, after=0)
            for r in p.runs:
                set_font(r, size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            borders(cells[i])
            cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[i].paragraphs:
                pspace(p, after=0, line=1.0)
                for r in p.runs:
                    set_font(r, size=font_size)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    doc.add_paragraph()
    return table


def add_code(doc, code, caption):
    p = doc.add_paragraph()
    pspace(p, after=2)
    r = p.add_run(caption)
    set_font(r, size=8.8, italic=True, color=GRAY)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "F3F4F6")
    borders(cell, "CBD5E1")
    cell_margins(cell, top=90, start=130, bottom=90, end=130)
    p = cell.paragraphs[0]
    pspace(p, after=0, line=1.0)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, size=8.0, name="Courier New", color="1F2937")
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def chapter(doc, title, subtitle=None):
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    borders(cell, BLUE, "8")
    cell_margins(cell, top=170, start=230, bottom=165, end=230)
    p = cell.paragraphs[0]
    p.style = "Heading 1"
    pspace(p, after=0)
    r = p.add_run(title)
    set_font(r, size=15.0, bold=True, color=NAVY)
    if subtitle:
        p = cell.add_paragraph()
        pspace(p, before=2, after=0)
        r = p.add_run(subtitle)
        set_font(r, size=9.3, italic=True, color=GRAY)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=20, after=6)
    r = p.add_run("Guide d’explication pour la soutenance")
    set_font(r, size=24, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, after=14)
    r = p.add_run("Projet ChatUser - RMI, XML-RPC, SOAP, REST et bonus TCP")
    set_font(r, size=14, bold=True, color=TEAL)

    add_callout(
        doc,
        "But de ce document",
        "Ce guide sert à préparer l’oral : il explique chaque travail pratique, le principe technique, ce qui a été réalisé dans le code et la différence avec la version précédente.",
        "oral",
    )
    add_table(
        doc,
        ["Élément", "Information"],
        [
            ("Projet", "ChatUser / TP 0"),
            ("Étudiants", "Salif Biaye, Ndeye Astou Diagouraga, Mouhamadou Tidiane Seck, Sountou Sakho"),
            ("Enseignant", "Pr. Ibrahima Fall"),
            ("Usage", "Support personnel pour expliquer clairement le projet pendant la présentation."),
        ],
        weights=[1.5, 4.5],
        header_fill=LIGHT_BLUE,
    )


def intro(doc):
    add_heading(doc, "1. Le fil rouge à retenir", 1)
    add_para(
        doc,
        "Le projet repose sur une idée simple : on garde la même application de chat, mais on change la manière dont le client et le serveur communiquent. C’est cette stabilité du besoin qui rend la comparaison intéressante.",
    )
    add_para(
        doc,
        "À l’oral, il faut donc toujours revenir à cette phrase : le métier est le même, mais le protocole change. Dans toutes les versions, un utilisateur rejoint une salle, envoie un message, reçoit les messages des autres et peut quitter la salle.",
    )
    add_callout(
        doc,
        "Phrase d’ouverture à dire",
        "Notre objectif n’était pas seulement de coder un chat. L’objectif était de comprendre plusieurs modèles de communication distribuée à travers le même cas pratique : RMI, XML-RPC, SOAP, REST et, en bonus, les sockets TCP.",
        "oral",
    )
    add_table(
        doc,
        ["Version", "Idée principale", "Ce qu’il faut retenir"],
        [
            ("RMI", "Appeler un objet Java distant", "Le serveur peut rappeler les clients grâce aux callbacks."),
            ("XML-RPC", "Appeler une procédure via HTTP/XML", "Le client interroge le serveur par polling."),
            ("SOAP", "Échanger des messages XML formels", "Le service est plus contractuel grâce au WSDL."),
            ("REST", "Exposer des endpoints HTTP/JSON", "C’est plus léger, plus lisible et proche des API modernes."),
            ("TCP", "Utiliser des sockets bas niveau", "Tout le protocole est géré manuellement."),
        ],
        weights=[1, 2.4, 3.2],
        header_fill=LIGHT_TEAL,
    )
    add_heading(doc, "2. Plan oral conseillé", 1)
    add_bullets(
        doc,
        [
            "Présenter le besoin commun : une application ChatUser client-serveur.",
            "Expliquer RMI comme point de départ orienté objet.",
            "Montrer pourquoi XML-RPC change le modèle : plus d’objet distant client, mais du polling.",
            "Expliquer SOAP comme version plus formelle et contractuelle de l’échange XML.",
            "Expliquer REST comme version plus moderne, plus légère et plus naturelle pour le web.",
            "Terminer par le bonus TCP : utile pour comprendre ce qui se passe sous les abstractions.",
            "Conclure avec le tableau de comparaison et les limites de chaque approche.",
        ],
    )


def rmi(doc):
    chapter(doc, "TP 1 : Java RMI", "Objets distants, registre RMI et callbacks")
    add_heading(doc, "Ce qu’il faut comprendre")
    add_para(
        doc,
        "RMI signifie Remote Method Invocation. L’idée est qu’un client Java peut appeler une méthode d’un objet qui se trouve sur le serveur comme si cet objet était local. En réalité, l’appel traverse le réseau, mais RMI cache une grande partie de cette complexité.",
    )
    add_para(
        doc,
        "Dans notre projet, l’objet distant principal est la salle de discussion. Le serveur crée un objet `ChatRoomImpl`, le publie dans le registre RMI, puis le client récupère cette référence distante pour appeler `subscribe`, `postMessage` et `unsubscribe`.",
    )
    add_callout(
        doc,
        "Particularité importante",
        "RMI permet aussi au serveur de rappeler le client. C’est pour cela qu’il existe une interface `ChatUser` : le client expose un objet distant, et le serveur appelle `displayMessage` pour pousser les messages.",
        "info",
    )
    add_code(
        doc,
        """
public interface ChatRoom extends Remote {
    void subscribe(ChatUser user, String pseudo) throws RemoteException;
    void unsubscribe(String pseudo) throws RemoteException;
    void postMessage(String pseudo, String message) throws RemoteException;
}
""",
        "Extrait à expliquer : le contrat distant de la salle RMI.",
    )
    add_heading(doc, "Comment l’expliquer à l’oral")
    add_bullets(
        doc,
        [
            "Le serveur publie une salle de discussion dans le registre RMI.",
            "Le client récupère cette salle avec un lookup.",
            "Quand le client envoie un message, il appelle une méthode distante.",
            "Le serveur garde la liste des utilisateurs connectés.",
            "Pour diffuser le message, le serveur rappelle chaque client avec `displayMessage`.",
        ],
    )
    add_heading(doc, "Différence avec une application locale")
    add_para(
        doc,
        "Dans une application locale, toutes les méthodes sont appelées dans la même JVM. Ici, les méthodes peuvent échouer à cause du réseau, donc elles déclarent `RemoteException`. C’est la preuve que l’objet paraît local dans le code, mais reste distant techniquement.",
    )
    add_callout(
        doc,
        "Phrase à retenir",
        "RMI est la version la plus orientée objet : le client manipule une référence distante, et le serveur peut pousser les messages grâce aux callbacks.",
        "oral",
    )


def xmlrpc(doc):
    chapter(doc, "TP 2 : XML-RPC", "HTTP + XML, appels de procédures et polling")
    add_heading(doc, "Ce qu’il faut comprendre")
    add_para(
        doc,
        "XML-RPC garde l’idée d’appeler une opération distante, mais il ne manipule plus des objets Java distants comme RMI. Le client envoie une requête HTTP contenant le nom de la méthode et les paramètres encodés en XML.",
    )
    add_para(
        doc,
        "Dans notre projet, le serveur démarre un `WebServer` sur le port 8080 et associe le nom `ChatRoom` à `ChatRoomImpl`. Les clients appellent des méthodes comme `ChatRoom.subscribe`, `ChatRoom.postMessage` et `ChatRoom.getMessages`.",
    )
    add_code(
        doc,
        """
PropertyHandlerMapping phm = new PropertyHandlerMapping();
phm.addHandler("ChatRoom", ChatRoomImpl.class);
xmlRpcServer.setHandlerMapping(phm);
""",
        "Extrait à expliquer : le serveur associe le nom ChatRoom à la classe métier.",
    )
    add_heading(doc, "Différence avec RMI")
    add_table(
        doc,
        ["Point", "RMI", "XML-RPC"],
        [
            ("Nature", "Objet distant Java", "Procédure distante appelée par nom"),
            ("Format", "Sérialisation Java", "XML dans une requête HTTP"),
            ("Réception", "Callback serveur vers client", "Polling avec `getMessages(lastIndex)`"),
            ("Interopérabilité", "Surtout Java", "Plus ouvert à d’autres langages"),
        ],
        weights=[1.4, 2.2, 2.6],
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "Le point clé : le polling")
    add_para(
        doc,
        "Avec XML-RPC, le serveur ne rappelle pas directement le client. Le client doit donc demander régulièrement : est-ce qu’il y a de nouveaux messages depuis mon dernier index ? C’est ce qu’on appelle le polling.",
    )
    add_callout(
        doc,
        "Phrase à retenir",
        "XML-RPC remplace le callback RMI par un échange requête/réponse : le client revient chercher les nouveaux messages avec `getMessages(lastIndex)`.",
        "oral",
    )


def soap(doc):
    chapter(doc, "TP 3 : SOAP", "Enveloppe XML, service web et WSDL")
    add_heading(doc, "Ce qu’il faut comprendre")
    add_para(
        doc,
        "SOAP est aussi basé sur XML, mais il est plus formel qu’XML-RPC. Un message SOAP est structuré dans une enveloppe : `Envelope`, éventuellement `Header`, puis `Body`. Le corps contient l’opération demandée.",
    )
    add_para(
        doc,
        "Dans notre projet, le serveur expose un service sur `/chat`. Une requête `GET /chat?wsdl` renvoie le contrat du service, tandis qu’une requête `POST` contient une enveloppe SOAP à parser.",
    )
    add_code(
        doc,
        """
Document document = SoapUtil.parse(request);
Element operation = SoapUtil.getSoapOperation(document);
String operationName = SoapUtil.localName(operation);
""",
        "Extrait à expliquer : le serveur parse l’enveloppe SOAP pour trouver l’opération.",
    )
    add_heading(doc, "Différence avec XML-RPC")
    add_table(
        doc,
        ["Point", "XML-RPC", "SOAP"],
        [
            ("Structure", "XML simple pour appeler une méthode", "Enveloppe SOAP plus formelle"),
            ("Contrat", "Pas de WSDL obligatoire", "WSDL pour décrire le service"),
            ("Lisibilité", "Plus simple", "Plus verbeux"),
            ("Usage", "RPC léger", "Services web contractuels"),
        ],
        weights=[1.4, 2.5, 2.5],
        header_fill=LIGHT_ORANGE,
    )
    add_heading(doc, "Le WSDL")
    add_para(
        doc,
        "Le WSDL sert à décrire le service : son nom, son adresse et les opérations disponibles. Dans une présentation, il faut dire que SOAP ne se contente pas d’envoyer du XML : il cherche aussi à formaliser le service.",
    )
    add_callout(
        doc,
        "Phrase à retenir",
        "SOAP est plus lourd qu’XML-RPC, mais il est plus contractuel : le WSDL permet de décrire officiellement le service.",
        "oral",
    )


def rest(doc):
    chapter(doc, "TP 4 : REST", "Endpoints HTTP, JSON et API web moderne")
    add_heading(doc, "Ce qu’il faut comprendre")
    add_para(
        doc,
        "REST change la manière de penser l’application. Au lieu d’appeler une procédure distante, on expose des endpoints HTTP. Chaque endpoint correspond à une action ou une ressource : s’inscrire, envoyer un message, récupérer les messages, se désinscrire.",
    )
    add_para(
        doc,
        "Dans notre projet, le serveur REST écoute sur le port 8082 et manipule du JSON. Le client envoie des requêtes `POST` ou `GET`, puis le serveur répond avec un statut JSON.",
    )
    add_table(
        doc,
        ["Endpoint", "Méthode", "Rôle"],
        [
            ("/chat/subscribe", "POST", "Inscrire un utilisateur"),
            ("/chat/message", "POST", "Publier un message"),
            ("/chat/messages?lastIndex=0", "GET", "Récupérer les nouveaux messages"),
            ("/chat/unsubscribe", "POST", "Quitter la salle"),
        ],
        weights=[2.6, 1.2, 3.2],
        header_fill=LIGHT_TEAL,
    )
    add_code(
        doc,
        """
server.createContext("/chat/subscribe", ChatServer::handleSubscribe);
server.createContext("/chat/message", ChatServer::handleMessage);
server.createContext("/chat/messages", ChatServer::handleMessages);
""",
        "Extrait à expliquer : chaque URL est reliée à une méthode de traitement.",
    )
    add_heading(doc, "Différence avec SOAP")
    add_table(
        doc,
        ["Point", "SOAP", "REST"],
        [
            ("Format", "XML enveloppé", "JSON plus léger"),
            ("Contrat", "WSDL", "Endpoints documentés"),
            ("Style", "Service web formel", "API web simple"),
            ("Test", "Moins direct sans outil SOAP", "Testable avec navigateur, curl ou Postman"),
        ],
        weights=[1.4, 2.5, 2.5],
        header_fill=LIGHT_GREEN,
    )
    add_callout(
        doc,
        "Phrase à retenir",
        "REST est la version la plus proche des API modernes : les actions sont exposées par des URL, les données sont en JSON, et les tests sont plus simples.",
        "oral",
    )


def tcp(doc):
    chapter(doc, "Bonus : sockets TCP", "Le réseau bas niveau derrière les abstractions")
    add_heading(doc, "Pourquoi c’est un bonus")
    add_para(
        doc,
        "La version TCP n’était pas le cœur demandé, mais elle est très utile pour comprendre ce que les technologies de plus haut niveau cachent. Avec TCP, on travaille directement avec `ServerSocket`, `Socket`, flux d’entrée/sortie et threads.",
    )
    add_code(
        doc,
        """
Socket socket = serverSocket.accept();
ClientHandler handler = new ClientHandler(socket);
new Thread(handler).start();
""",
        "Extrait à expliquer : un thread est créé pour chaque client connecté.",
    )
    add_heading(doc, "Différence avec REST, SOAP et XML-RPC")
    add_para(
        doc,
        "REST, SOAP et XML-RPC reposent sur HTTP et imposent déjà une certaine structure. TCP, lui, ne définit pas le protocole applicatif. C’est à nous de décider que la première ligne correspond au pseudo et que chaque ligne suivante est un message.",
    )
    add_table(
        doc,
        ["Aspect", "Versions HTTP/RPC", "TCP sockets"],
        [
            ("Structure", "Déjà définie par le protocole ou le format", "À définir nous-mêmes"),
            ("Messages", "XML ou JSON", "Lignes de texte"),
            ("Concurrence", "Gérée en partie par le serveur HTTP", "Thread par client dans notre code"),
            ("Avantage", "Plus simple à intégrer", "Permet de comprendre le réseau bas niveau"),
            ("Limite", "Moins proche du bas niveau", "Plus manuel et plus fragile"),
        ],
        weights=[1.4, 2.8, 2.8],
        header_fill=LIGHT_ORANGE,
    )
    add_callout(
        doc,
        "Phrase à retenir",
        "TCP montre la mécanique de base : connexion, flux, thread et diffusion. Les autres technologies ajoutent une structure au-dessus de cette communication réseau.",
        "oral",
    )


def ant_fop(doc):
    chapter(doc, "Annexes à expliquer si on te demande", "Ant et FOP sans te perdre")
    add_heading(doc, "Apache Ant")
    add_para(
        doc,
        "Ant n’est pas une technologie de communication. C’est un outil d’automatisation. Il sert à compiler, générer la documentation, créer les archives et lancer les programmes avec des commandes reproductibles.",
    )
    add_table(
        doc,
        ["Cible", "Explication simple à donner"],
        [
            ("init", "Prépare les dossiers de sortie."),
            ("compile", "Compile les fichiers Java."),
            ("doc", "Génère la documentation Javadoc."),
            ("arch", "Crée une archive JAR."),
            ("run-server", "Lance le serveur."),
            ("run-client", "Lance un client."),
            ("all", "Exécute tout le build."),
        ],
        weights=[1.4, 4.6],
        header_fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "Phrase à dire",
        "Ant nous permet de rendre les tests reproductibles : au lieu de taper toutes les commandes Java à la main, on lance des cibles définies dans `build.xml`.",
        "oral",
    )
    add_heading(doc, "Apache FOP")
    add_para(
        doc,
        "FOP sert à générer des PDF à partir de documents XSL-FO. Le principe est de séparer les données, la mise en forme et la sortie finale. Dans le projet, les données sont dans XML, la mise en page est décrite en XSL-FO, puis FOP produit le PDF.",
    )
    add_table(
        doc,
        ["Élément", "Rôle"],
        [
            ("XML", "Contient les données."),
            ("XSL / XSL-FO", "Transforme et décrit la mise en page."),
            ("Fo2Pdf.java", "Pilote FOP depuis Java."),
            ("PDF", "Résultat final généré."),
        ],
        weights=[1.6, 4.4],
        header_fill=LIGHT_TEAL,
    )
    add_callout(
        doc,
        "Phrase à dire",
        "FOP montre une autre chaîne technique : on ne communique pas entre client et serveur, mais on transforme des données structurées en document final.",
        "oral",
    )


def comparisons(doc):
    chapter(doc, "Comparaison finale à maîtriser", "La slide mentale qui sauve la soutenance")
    add_heading(doc, "Tableau ultra important")
    add_table(
        doc,
        ["Technologie", "Question à se poser", "Réponse courte"],
        [
            ("RMI", "Est-ce que je manipule un objet distant ?", "Oui, et le serveur peut rappeler le client."),
            ("XML-RPC", "Est-ce que j’appelle une méthode par HTTP/XML ?", "Oui, mais sans callback : polling."),
            ("SOAP", "Est-ce que l’échange est formel et contractuel ?", "Oui, grâce à l’enveloppe SOAP et au WSDL."),
            ("REST", "Est-ce que j’expose des endpoints HTTP/JSON ?", "Oui, c’est le style API web moderne."),
            ("TCP", "Est-ce que je gère moi-même le protocole ?", "Oui, ligne par ligne avec sockets et threads."),
        ],
        weights=[1.2, 2.7, 3.1],
        header_fill=LIGHT_BLUE,
    )
    add_heading(doc, "Différences en chaîne")
    add_bullets(
        doc,
        [
            "RMI vers XML-RPC : on quitte l’objet distant Java pour un appel distant HTTP/XML.",
            "XML-RPC vers SOAP : on passe d’un RPC simple à un service plus formel avec enveloppe et WSDL.",
            "SOAP vers REST : on abandonne la lourdeur XML/WSDL pour des endpoints HTTP et du JSON.",
            "REST vers TCP : on descend au niveau réseau, sans structure imposée par HTTP ou XML/JSON.",
        ],
    )
    add_heading(doc, "Questions possibles du professeur")
    add_table(
        doc,
        ["Question", "Réponse claire"],
        [
            ("Pourquoi garder le même chat dans toutes les versions ?", "Pour comparer les technologies sans changer le besoin fonctionnel."),
            ("Pourquoi RMI a des callbacks et pas XML-RPC ?", "Parce qu’en RMI le client expose aussi un objet distant `ChatUser`; XML-RPC reste en requête/réponse."),
            ("Pourquoi utiliser `lastIndex` ?", "Pour récupérer seulement les messages non lus et éviter de renvoyer tout l’historique."),
            ("Pourquoi SOAP est plus lourd ?", "Parce qu’il ajoute une enveloppe XML et un contrat WSDL."),
            ("Pourquoi REST est plus moderne ?", "Parce qu’il utilise HTTP naturellement, des endpoints lisibles et JSON."),
            ("Pourquoi TCP est plus manuel ?", "Parce qu’il faut gérer soi-même les lignes, les connexions, les threads et la diffusion."),
        ],
        weights=[2.2, 4.8],
        header_fill=LIGHT_GREEN,
        font_size=8.4,
    )
    add_callout(
        doc,
        "Conclusion orale possible",
        "Ce projet m’a permis de comprendre qu’une application distribuée ne dépend pas seulement de son code métier. Le choix du protocole change la structure du code, la manière de tester, l’interopérabilité et la complexité de maintenance.",
        "ok",
    )


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    cover(doc)
    intro(doc)
    rmi(doc)
    xmlrpc(doc)
    soap(doc)
    rest(doc)
    tcp(doc)
    ant_fop(doc)
    comparisons(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
