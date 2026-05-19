from pathlib import Path
import math
import random
import sys
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\DELL\Downloads\i-fall")
OUT = ROOT / "rapport-i-fall-complet.docx"
FIG_DIR = ROOT / "figures_croquis_ifall"
LOGO_UCAD = Path(r"C:\Users\DELL\Downloads\logo_ucad.png")
REPO_URL = "https://github.com/salifbiaye/projet-tp-0"

SKILL_SCRIPTS = Path(
    r"C:\Users\DELL\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
sys.path.append(str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights

CONTENT_WIDTH_DXA = 10080

NAVY = "0B1F3A"
TEAL = "0F7182"
BLUE = "2563EB"
GREEN = "0F766E"
ORANGE = "EA580C"
RED = "B91C1C"
GRAY = "64748B"
TEXT = "111827"
MUTED = "334155"
BORDER = "CBD5E1"
LIGHT_BLUE = "EFF6FF"
LIGHT_TEAL = "E6FFFA"
LIGHT_GREEN = "ECFDF5"
LIGHT_ORANGE = "FFF7ED"
LIGHT_GRAY = "F8FAFC"
LIGHT_RED = "FEF2F2"
BOOKMARK_ID = 0

ANCHORS = {
    "cover": "cover",
    "summary": "summary",
    "figures": "figures",
    "ch1": "ch1_intro",
    "ch2": "ch2_rmi",
    "ch3": "ch3_xmlrpc",
    "ch4": "ch4_soap",
    "ch5": "ch5_rest",
    "ch6": "ch6_comparison",
    "ch7": "ch7_validation",
    "bonus_tcp": "bonus_tcp",
    "annexes": "annexes",
}


def fig_anchor(number):
    return "fig_" + number.replace(".", "_").replace("-", "_")

FIG_TITLES = {
    "1.1": "Architecture commune du projet ChatUser",
    "2.1": "Architecture Java RMI avec registre et callbacks",
    "3.1": "Communication XML-RPC avec appel HTTP et polling",
    "4.1": "Traitement SOAP avec enveloppe XML et WSDL",
    "5.1": "Architecture REST avec endpoints HTTP et JSON",
    "6.1": "Comparaison des modèles de communication distribuée",
    "7.1": "Chaîne de validation commune des applications ChatUser",
    "B.1": "Bonus technique : variante TCP sockets",
    "A.1": "Pipeline Apache Ant du projet",
    "A.2": "Pipeline Apache FOP de génération PDF",
}


def font_path(*names):
    fonts_dir = Path(r"C:\Windows\Fonts")
    for name in names:
        candidate = fonts_dir / name
        if candidate.exists():
            return str(candidate)
    return None


FONT_BODY = font_path("arial.ttf", "calibri.ttf")
FONT_BOLD = font_path("arialbd.ttf", "calibrib.ttf")
FONT_HAND = font_path("comic.ttf", "segoepr.ttf", "arial.ttf")


def pil_font(size=28, bold=False, hand=False):
    path = FONT_HAND if hand else (FONT_BOLD if bold else FONT_BODY)
    try:
        return ImageFont.truetype(path, size=size)
    except Exception:
        return ImageFont.load_default()


def set_font(run, size=10.5, bold=False, italic=False, color=TEXT, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def pspace(paragraph, before=0, after=3, line=1.02):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcBorders")
    if node is None:
        node = OxmlElement("w:tcBorders")
        tc_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_node = node.find(qn(f"w:{edge}"))
        if edge_node is None:
            edge_node = OxmlElement(f"w:{edge}")
            node.append(edge_node)
        edge_node.set(qn("w:val"), "single")
        edge_node.set(qn("w:sz"), size)
        edge_node.set(qn("w:space"), "0")
        edge_node.set(qn("w:color"), color)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = tc_mar.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def add_field(paragraph, instr):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(text)
    run._r.append(separate)
    run._r.append(fallback)
    run._r.append(end)


def add_bookmark(paragraph, name):
    global BOOKMARK_ID
    BOOKMARK_ID += 1
    bookmark_id = str(BOOKMARK_ID)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text, anchor, size=9.2):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Arial")
    font.set(qn("w:hAnsi"), "Arial")
    r_pr.append(color)
    r_pr.append(underline)
    r_pr.append(sz)
    r_pr.append(font)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    for r in paragraph.runs:
        set_font(r, size=size, color=BLUE)


def add_external_hyperlink(paragraph, text, url, size=9.5):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Arial")
    font.set(qn("w:hAnsi"), "Arial")
    r_pr.append(color)
    r_pr.append(underline)
    r_pr.append(sz)
    r_pr.append(font)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    for r in paragraph.runs:
        set_font(r, size=size, color=BLUE)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.02

    for style_name, size, color in [
        ("Title", 22, NAVY),
        ("Subtitle", 11, GRAY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11.5, BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(8 if style_name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(4)

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.text = ""
        r = header.add_run("Rapport de travaux pratiques - Systèmes distribués et services web")
        set_font(r, size=8.6, bold=True, color=GRAY)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run("Page ")
        set_font(r, size=8.5, color=GRAY)
        add_field(footer, "PAGE")
        r = footer.add_run(" / ")
        set_font(r, size=8.5, color=GRAY)
        add_field(footer, "NUMPAGES")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, value, font_size=8.7, bold=False, color=TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    pspace(p, after=0, line=1.0)
    if isinstance(value, tuple) and len(value) == 2:
        label, anchor = value
        add_internal_hyperlink(p, str(label), str(anchor), size=font_size)
    else:
        for idx, line in enumerate(str(value).split("\n")):
            if idx:
                p.add_run("\n")
            r = p.add_run(line)
            set_font(r, size=font_size, bold=bold, color=color)


def add_table(doc, headers, rows, weights=None, header_fill=LIGHT_GRAY, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if weights is None:
        weights = [1] * len(headers)
    widths = column_widths_from_weights(weights, CONTENT_WIDTH_DXA)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, header_fill)
        borders(cell)
        cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        write_cell(cell, header, font_size=font_size, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            borders(cells[i])
            cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            write_cell(cells[i], value, font_size=font_size, color=TEXT)

    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    doc.add_paragraph()
    return table


def add_paragraph(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    pspace(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        pspace(p, after=2, line=1.02)
        r = p.add_run(item)
        set_font(r, size=10.1)


def add_callout(doc, title, body, kind="info"):
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "method": (LIGHT_TEAL, TEAL),
        "warn": (LIGHT_ORANGE, ORANGE),
        "ok": (LIGHT_GREEN, GREEN),
        "risk": (LIGHT_RED, RED),
    }
    fill, accent = colors.get(kind, colors["info"])
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell, color=accent, size="8")
    cell_margins(cell, top=110, start=150, bottom=110, end=150)
    p = cell.paragraphs[0]
    pspace(p, after=2)
    r = p.add_run(title)
    set_font(r, size=10.2, bold=True, color=accent)
    p = cell.add_paragraph()
    pspace(p, after=0, line=1.02)
    r = p.add_run(body)
    set_font(r, size=9.6, color=TEXT)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def add_code_block(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        pspace(p, before=2, after=2)
        r = p.add_run(caption)
        set_font(r, size=9.2, italic=True, color=MUTED)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, "F3F4F6")
    borders(cell, color="CBD5E1", size="4")
    cell_margins(cell, top=100, start=130, bottom=100, end=130)
    p = cell.paragraphs[0]
    pspace(p, after=0, line=1.0)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, size=8.4, color="1F2937", name="Courier New")
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def add_capture_placeholder(doc, title, instruction):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, "FFFFFF")
    borders(cell, color=TEAL, size="6")
    cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    pspace(p, after=3)
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=TEAL)
    p = cell.add_paragraph()
    pspace(p, after=0, line=1.0)
    r = p.add_run(instruction)
    set_font(r, size=9.0, italic=True, color=GRAY)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=10, after=6)
    r = p.add_run("[ Zone réservée à la capture d’écran ]")
    set_font(r, size=10.2, bold=True, color=MUTED)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def chapter(doc, title, subtitle=None, anchor=None, page_break=True):
    if page_break and len(doc.paragraphs) > 1:
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    borders(cell, color=BLUE, size="8")
    cell_margins(cell, top=130, start=170, bottom=125, end=170)
    p = cell.paragraphs[0]
    p.style = "Heading 1"
    pspace(p, after=0)
    if anchor:
        add_bookmark(p, anchor)
    r = p.add_run(title)
    set_font(r, size=15.0, bold=True, color=NAVY)
    if subtitle:
        p = cell.add_paragraph()
        pspace(p, after=0)
        r = p.add_run(subtitle)
        set_font(r, size=9.3, italic=True, color=MUTED)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=0)
    doc.add_paragraph()


def section_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    pspace(p, before=6 if level == 2 else 4, after=3)
    return p


def wrap_text(draw, text, font, max_width):
    lines = []
    for paragraph in text.split("\n"):
        current = ""
        for word in paragraph.split():
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def hand_line(draw, points, fill=(35, 35, 35), width=3, jitter=1):
    random.seed(sum(int(v) for p in points for v in p) + width)
    for _ in range(2):
        jittered = [(x + random.randint(-jitter, jitter), y + random.randint(-jitter, jitter)) for x, y in points]
        draw.line(jittered, fill=fill, width=width)


def draw_box(draw, xy, title, body="", fill=(255, 255, 255), hatch=False, title_color=(15, 113, 130)):
    x1, y1, x2, y2 = xy
    if fill:
        draw.rectangle(xy, fill=fill)
    if hatch:
        for x in range(x1 - 20, x2 + 20, 16):
            draw.line((x, y2, x + 90, y1), fill=(224, 224, 224), width=2)
    hand_line(draw, [(x1, y1), (x2, y1), (x2, y2), (x1, y2), (x1, y1)], width=3)
    font_t = pil_font(23, bold=True, hand=True)
    font_b = pil_font(18, hand=True)
    draw.text((x1 + 16, y1 + 14), title, font=font_t, fill=title_color)
    if body:
        y = y1 + 50
        for line in wrap_text(draw, body, font_b, (x2 - x1) - 30):
            draw.text((x1 + 16, y), line, font=font_b, fill=(25, 25, 25))
            y += 25


def arrow(draw, start, end, dashed=False, label=None):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        total = math.hypot(x2 - x1, y2 - y1)
        steps = max(1, int(total // 18))
        for i in range(0, steps, 2):
            sx = x1 + (x2 - x1) * i / steps
            sy = y1 + (y2 - y1) * i / steps
            ex = x1 + (x2 - x1) * min(i + 1, steps) / steps
            ey = y1 + (y2 - y1) * min(i + 1, steps) / steps
            draw.line((sx, sy, ex, ey), fill=(40, 40, 40), width=3)
    else:
        hand_line(draw, [(x1, y1), (x2, y2)], width=3)
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 12
    p1 = (x2 - size * math.cos(angle - 0.5), y2 - size * math.sin(angle - 0.5))
    p2 = (x2 - size * math.cos(angle + 0.5), y2 - size * math.sin(angle + 0.5))
    draw.polygon([(x2, y2), p1, p2], fill=(40, 40, 40))
    if label:
        font_l = pil_font(17, hand=True)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.text((mx - 70, my - 28), label, font=font_l, fill=(15, 113, 130))


def make_canvas(title):
    img = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((24, 24, 1476, 736), outline=(25, 25, 25), width=3)
    font = pil_font(34, bold=True, hand=True)
    draw.text((70, 50), title, font=font, fill=(15, 113, 130))
    hand_line(draw, [(70, 100), (1390, 98)], fill=(120, 120, 120), width=2)
    return img, draw


def save_fig(name, builder):
    FIG_DIR.mkdir(exist_ok=True)
    path = FIG_DIR / f"{name}.png"
    img, draw = make_canvas(FIG_TITLES[name])
    builder(draw)
    img.save(path, quality=95)
    return path


def build_figures():
    figures = {}

    def common(draw):
        draw_box(draw, (80, 170, 430, 300), "Clients Swing", "Alice, Bob, Charlie\nInterface utilisateur", fill=(248, 250, 252))
        draw_box(draw, (570, 160, 930, 315), "Serveur ChatRoom", "Inscription\nMessages\nDiffusion / consultation", fill=(236, 253, 245), hatch=True)
        draw_box(draw, (1080, 170, 1400, 300), "Etat partagé", "Utilisateurs\nHistorique\nPseudo unique", fill=(255, 247, 237))
        draw_box(draw, (210, 455, 1290, 630), "Technologies étudiées", "RMI : objet distant et callback    |    XML-RPC : HTTP + XML\nSOAP : enveloppe XML + WSDL        |    REST : endpoints HTTP + JSON", fill=(255, 255, 255))
        arrow(draw, (430, 235), (570, 235), label="requête")
        arrow(draw, (930, 235), (1080, 235), label="mise à jour")
        arrow(draw, (750, 315), (750, 455), dashed=True, label="même besoin")
        draw.text((545, 690), "Un même cas fonctionnel sert à comparer plusieurs styles de communication.", font=pil_font(20, hand=True), fill=(20, 20, 20))

    figures["1.1"] = save_fig("1.1", common)

    def rmi(draw):
        draw_box(draw, (80, 190, 360, 350), "Client Alice", "ChatClient\nChatUser distant", fill=(248, 250, 252))
        draw_box(draw, (550, 170, 880, 335), "RMI Registry", "Nom publié :\nChatRoom", fill=(239, 246, 255), hatch=True)
        draw_box(draw, (1030, 190, 1380, 350), "Serveur", "ChatRoomImpl\nobjet distant", fill=(236, 253, 245))
        draw_box(draw, (580, 520, 920, 670), "Client Bob", "displayMessage()\nappelé par le serveur", fill=(255, 247, 237))
        arrow(draw, (360, 245), (550, 245), label="lookup")
        arrow(draw, (880, 245), (1030, 245), label="référence")
        arrow(draw, (1030, 310), (760, 520), dashed=True, label="callback")
        arrow(draw, (360, 305), (1030, 305), dashed=True, label="postMessage")
        draw.text((100, 705), "RMI garde une logique objet : le client appelle une méthode distante et le serveur peut rappeler les clients inscrits.", font=pil_font(20, hand=True), fill=(20, 20, 20))

    figures["2.1"] = save_fig("2.1", rmi)

    def xmlrpc(draw):
        draw_box(draw, (85, 185, 390, 340), "Client", "XmlRpcClient\nméthodes distantes", fill=(248, 250, 252))
        draw_box(draw, (570, 170, 910, 350), "HTTP + XML", "<methodCall>\nChatRoom.postMessage", fill=(239, 246, 255), hatch=True)
        draw_box(draw, (1090, 185, 1400, 340), "Serveur", "WebServer 8080\nHandler ChatRoom", fill=(236, 253, 245))
        draw_box(draw, (620, 520, 980, 665), "Liste messages", "lastIndex\nnouveaux messages", fill=(255, 247, 237))
        arrow(draw, (390, 245), (570, 245), label="appel")
        arrow(draw, (910, 245), (1090, 245), label="exécute")
        arrow(draw, (1240, 340), (980, 520), dashed=True, label="stocke")
        arrow(draw, (980, 590), (390, 315), dashed=True, label="polling")
        draw.text((90, 705), "Le client ne reçoit pas de callback : il interroge régulièrement le serveur avec getMessages(lastIndex).", font=pil_font(20, hand=True), fill=(20, 20, 20))

    figures["3.1"] = save_fig("3.1", xmlrpc)

    def soap(draw):
        draw_box(draw, (80, 180, 370, 335), "Client SOAP", "Requête POST\nopération XML", fill=(248, 250, 252))
        draw_box(draw, (535, 160, 900, 355), "Enveloppe SOAP", "Envelope\nHeader optionnel\nBody obligatoire", fill=(239, 246, 255), hatch=True)
        draw_box(draw, (1060, 180, 1400, 335), "Service /chat", "SoapUtil.parse()\nChatRoomImpl", fill=(236, 253, 245))
        draw_box(draw, (590, 510, 950, 660), "WSDL", "GET /chat?wsdl\ncontrat du service", fill=(255, 247, 237))
        arrow(draw, (370, 245), (535, 245), label="HTTP")
        arrow(draw, (900, 245), (1060, 245), label="parse")
        arrow(draw, (1220, 335), (950, 510), dashed=True, label="décrit")
        arrow(draw, (1060, 300), (370, 300), dashed=True, label="SOAP response")
        draw.text((90, 705), "SOAP formalise l’échange : le corps XML identifie l’opération et la réponse reprend une structure enveloppée.", font=pil_font(20, hand=True), fill=(20, 20, 20))

    figures["4.1"] = save_fig("4.1", soap)

    def rest(draw):
        draw_box(draw, (80, 165, 390, 340), "Client", "Swing, navigateur\nou curl", fill=(248, 250, 252))
        draw_box(draw, (535, 145, 940, 360), "Endpoints HTTP", "POST /subscribe\nPOST /message\nGET /messages", fill=(239, 246, 255), hatch=True)
        draw_box(draw, (1085, 165, 1400, 340), "Serveur REST", "HttpServer 8082\nJsonUtil", fill=(236, 253, 245))
        draw_box(draw, (580, 520, 1020, 665), "Réponse JSON", "{\"status\":\"OK\"}\nmessages + nextIndex", fill=(255, 247, 237))
        arrow(draw, (390, 250), (535, 250), label="JSON")
        arrow(draw, (940, 250), (1085, 250), label="route")
        arrow(draw, (1120, 340), (1020, 545), dashed=True, label="format")
        arrow(draw, (580, 590), (390, 315), dashed=True, label="réponse")
        draw.text((90, 705), "REST expose des ressources par URL et s’appuie sur des méthodes HTTP lisibles pour piloter le chat.", font=pil_font(20, hand=True), fill=(20, 20, 20))

    figures["5.1"] = save_fig("5.1", rest)

    def comparison(draw):
        draw_box(draw, (80, 165, 390, 340), "RMI", "Objet distant\nCallback\nJava uniquement", fill=(236, 253, 245))
        draw_box(draw, (435, 165, 735, 340), "XML-RPC", "RPC simple\nXML sur HTTP\nPolling", fill=(239, 246, 255))
        draw_box(draw, (780, 165, 1075, 340), "SOAP", "Contrat WSDL\nEnvelope XML\nService formel", fill=(255, 247, 237))
        draw_box(draw, (1120, 165, 1410, 340), "REST", "Endpoints\nJSON\nInteropérable", fill=(248, 250, 252))
        draw_box(draw, (190, 500, 1320, 650), "Lecture comparative", "Plus on va vers REST, plus l’échange devient web et interopérable.\nPlus on va vers RMI, plus l’échange ressemble à une application Java objet.", fill=(255, 255, 255), hatch=True)
        arrow(draw, (230, 340), (330, 500), dashed=True)
        arrow(draw, (585, 340), (645, 500), dashed=True)
        arrow(draw, (930, 340), (930, 500), dashed=True)
        arrow(draw, (1265, 340), (1190, 500), dashed=True)

    figures["6.1"] = save_fig("6.1", comparison)

    def validation(draw):
        steps = [
            ((80, 200, 300, 330), "1. Compiler", "ant compile"),
            ((360, 200, 580, 330), "2. Serveur", "ant run-server"),
            ((640, 200, 860, 330), "3. Clients", "ant run-client"),
            ((920, 200, 1140, 330), "4. Tester", "messages"),
            ((1200, 200, 1420, 330), "5. Archiver", "ant all"),
        ]
        for xy, title, body in steps:
            draw_box(draw, xy, title, body, fill=(248, 250, 252))
        for i in range(len(steps) - 1):
            arrow(draw, (steps[i][0][2], 265), (steps[i + 1][0][0], 265), label="")
        draw_box(draw, (270, 510, 1240, 660), "Résultat attendu", "Deux clients connectés voient les notifications d’entrée, les messages envoyés et les notifications de sortie.", fill=(236, 253, 245), hatch=True)
        arrow(draw, (750, 330), (750, 510), dashed=True)

    figures["7.1"] = save_fig("7.1", validation)

    def tcp(draw):
        draw_box(draw, (80, 180, 390, 340), "Clients", "Socket\npseudo puis lignes", fill=(248, 250, 252))
        draw_box(draw, (560, 165, 920, 355), "Serveur TCP", "ServerSocket 8083\naccept()", fill=(239, 246, 255), hatch=True)
        draw_box(draw, (1080, 180, 1400, 340), "Threads", "ClientHandler\nun thread/client", fill=(236, 253, 245))
        draw_box(draw, (590, 520, 980, 665), "Broadcast", "ChatRoomImpl\nsendMessage()", fill=(255, 247, 237))
        arrow(draw, (390, 250), (560, 250), label="connexion")
        arrow(draw, (920, 250), (1080, 250), label="création")
        arrow(draw, (1240, 340), (980, 520), dashed=True, label="diffusion")
        arrow(draw, (590, 590), (390, 310), dashed=True, label="retour")
        draw.text((95, 705), "Variante utile pour comprendre le réseau bas niveau, mais placée en bonus car elle n’était pas le cœur demandé.", font=pil_font(20, hand=True), fill=(20, 20, 20))

    figures["B.1"] = save_fig("B.1", tcp)

    def ant(draw):
        steps = [
            ((100, 210, 340, 340), "src/", "Bonjour.java"),
            ((420, 210, 660, 340), "compile", "javac vers bin/"),
            ((740, 210, 980, 340), "doc", "Javadoc"),
            ((1060, 210, 1320, 340), "arch", "ProjetAnt.jar"),
        ]
        for xy, title, body in steps:
            draw_box(draw, xy, title, body, fill=(248, 250, 252))
        for i in range(len(steps) - 1):
            arrow(draw, (steps[i][0][2], 275), (steps[i + 1][0][0], 275), label="")
        draw_box(draw, (330, 520, 1180, 665), "Cible all", "clean -> init -> compile -> doc -> arch\nAutomatisation reproductible du projet Java", fill=(236, 253, 245), hatch=True)
        arrow(draw, (740, 340), (750, 520), dashed=True)

    figures["A.1"] = save_fig("A.1", ant)

    def fop(draw):
        draw_box(draw, (80, 190, 350, 340), "XML", "equipe.xml\ndonnées", fill=(248, 250, 252))
        draw_box(draw, (480, 190, 760, 340), "XSL / XSL-FO", "statistiques-fo.xsl\nmise en page", fill=(239, 246, 255), hatch=True)
        draw_box(draw, (900, 190, 1165, 340), "Apache FOP", "Fo2Pdf.java\nDriver FOP", fill=(236, 253, 245))
        draw_box(draw, (1250, 190, 1420, 340), "PDF", "sortie\nimprimable", fill=(255, 247, 237))
        arrow(draw, (350, 265), (480, 265), label="transforme")
        arrow(draw, (760, 265), (900, 265), label="rend")
        arrow(draw, (1165, 265), (1250, 265), label="")
        draw_box(draw, (260, 515, 1240, 660), "Idée clé", "Le contenu XML reste séparé de la présentation ; XSL-FO décrit la page et FOP produit le document final.", fill=(255, 255, 255), hatch=True)

    figures["A.2"] = save_fig("A.2", fop)

    return figures


def add_figure(doc, figures, number):
    title = FIG_TITLES[number]
    p = doc.add_paragraph()
    pspace(p, before=2, after=3)
    r = p.add_run(f"La figure {number} présente {title.lower()}.")
    set_font(r, size=9.8, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, after=2)
    run = p.add_run()
    run.add_picture(str(figures[number]), width=Cm(14.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, after=5)
    add_bookmark(p, fig_anchor(number))
    r = p.add_run(f"Figure {number} : {title}")
    set_font(r, size=8.7, italic=True, color=MUTED)


def add_cover(doc):
    if LOGO_UCAD.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_bookmark(p, ANCHORS["cover"])
        run = p.add_run()
        run.add_picture(str(LOGO_UCAD), width=Cm(2.4))
    else:
        p = doc.add_paragraph()
        add_bookmark(p, ANCHORS["cover"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=6, after=3)
    r = p.add_run("Université Cheikh Anta Diop de Dakar")
    set_font(r, size=13, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, after=18)
    r = p.add_run("Systèmes distribués et services web")
    set_font(r, size=11, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=8, after=8)
    r = p.add_run("Rapport de travaux pratiques")
    set_font(r, size=24, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, after=18)
    r = p.add_run("Projet i-fall / ChatUser")
    set_font(r, size=16, bold=True, color=TEAL)

    add_callout(
        doc,
        "Objet du document",
        "Présenter de manière claire, structurée et détaillée les réalisations autour de l’application ChatUser : RMI, XML-RPC, SOAP, REST, comparaison des architectures, guide d’exécution, puis annexes Ant et FOP.",
        kind="method",
    )

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Étudiants", "Salif Biaye\nNdeye Astou Diagouraga\nMouhamadou Tidiane Seck\nSountou Sakho"),
        ("Enseignant", "Pr. Ibrahima Fall"),
        ("Nature du livrable", "Rapport Word complet, sans LaTeX"),
        ("Technologies étudiées", "Java RMI, XML-RPC, SOAP, REST, Ant, Apache FOP"),
        ("Code source", REPO_URL),
        ("Date", "Mai 2026"),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        for j in range(2):
            cell = table.cell(i, j)
            borders(cell)
            cell_margins(cell, top=95, bottom=95)
            if j == 0:
                shade(cell, LIGHT_GRAY)
            for p in cell.paragraphs:
                pspace(p, after=0)
                for r in p.runs:
                    set_font(r, size=9.7, bold=(j == 0), color=NAVY if j == 0 else TEXT)
    apply_table_geometry(
        table,
        column_widths_from_weights([1.7, 4.3], CONTENT_WIDTH_DXA),
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=0,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=22, after=0)
    r = p.add_run("Version académique structurée avec schémas explicatifs et extraits de code commentés")
    set_font(r, size=9.4, italic=True, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=6, after=0)
    r = p.add_run("Code source : ")
    set_font(r, size=8.8, color=GRAY)
    add_external_hyperlink(p, REPO_URL, REPO_URL, size=8.8)


def add_front_matter(doc):
    chapter(doc, "Sommaire", "Plan manuel cliquable du document", anchor=ANCHORS["summary"])
    rows = [
        (("Introduction générale aux systèmes distribués", ANCHORS["ch1"]), "Lien"),
        (("Application ChatUser avec Java RMI", ANCHORS["ch2"]), "Lien"),
        (("Application ChatUser avec XML-RPC", ANCHORS["ch3"]), "Lien"),
        (("Application ChatUser avec SOAP", ANCHORS["ch4"]), "Lien"),
        (("Application ChatUser avec REST", ANCHORS["ch5"]), "Lien"),
        (("Comparaison générale des architectures", ANCHORS["ch6"]), "Lien"),
        (("Guide d’exécution, validation et captures", ANCHORS["ch7"]), "Lien"),
        (("Bonus technique : variante TCP sockets", ANCHORS["bonus_tcp"]), "Lien"),
        (("Annexes : Projet Ant, Apache FOP et références", ANCHORS["annexes"]), "Lien"),
    ]
    add_table(doc, ["Partie", "Accès"], rows, weights=[5, 1], header_fill=LIGHT_BLUE)

    p = section_heading(doc, "Table des figures", 2)
    add_bookmark(p, ANCHORS["figures"])
    fig_rows = [((f"Figure {num}", fig_anchor(num)), title) for num, title in FIG_TITLES.items()]
    add_table(doc, ["Figure", "Titre"], fig_rows, weights=[1.2, 4.8], header_fill=LIGHT_TEAL)


def chapter_intro(doc, figures):
    chapter(doc, "Chapitre 1 : Introduction générale aux systèmes distribués", "Comprendre le contexte avant les implémentations", anchor=ANCHORS["ch1"])
    section_heading(doc, "1.1 Contexte")
    add_paragraph(
        doc,
        "Le projet i-fall porte sur une application de discussion nommée ChatUser. Le besoin fonctionnel reste volontairement simple : plusieurs utilisateurs doivent pouvoir rejoindre une salle de discussion, envoyer des messages et recevoir les messages des autres participants. Cette simplicité est utile pédagogiquement, car elle permet de comparer plusieurs technologies distribuées sans changer le métier de l’application.",
    )
    add_paragraph(
        doc,
        "Dans un système distribué, les composants ne s’exécutent pas forcément dans la même machine virtuelle, ni même sur la même machine physique. Il faut donc définir un moyen de communication entre client et serveur : appel d’objet distant, appel de procédure distante, échange HTTP, message XML, message JSON ou socket réseau. Le projet illustre ces approches à travers des versions Java différentes.",
    )
    add_callout(
        doc,
        "Principe directeur",
        "Le rapport ne se contente pas de décrire les fichiers : il explique la logique de chaque architecture, les choix de communication et la manière de tester le comportement attendu.",
        "method",
    )
    section_heading(doc, "1.2 Objectif fonctionnel commun")
    add_paragraph(
        doc,
        "Toutes les variantes reposent sur le même scénario : un serveur maintient l’état de la salle, un client saisit un pseudo, le serveur accepte ou refuse l’inscription, puis les messages sont diffusés ou récupérés selon la technologie employée. La comparaison est donc fiable, parce que le besoin reste stable alors que le mécanisme réseau change.",
    )
    add_bullets(
        doc,
        [
            "Inscrire un utilisateur avec un pseudo unique.",
            "Envoyer un message depuis un client vers la salle commune.",
            "Mettre les nouveaux messages à disposition des autres clients.",
            "Gérer la sortie d’un utilisateur et notifier les autres participants.",
            "Compiler, exécuter et archiver les projets avec Apache Ant.",
        ],
    )
    add_figure(doc, figures, "1.1")
    section_heading(doc, "1.3 Périmètre retenu")
    add_paragraph(
        doc,
        "Le cœur du rapport traite les versions RMI, XML-RPC, SOAP et REST. La variante TCP sockets est conservée en section bonus, car elle aide à comprendre les mécanismes bas niveau mais ne constitue pas le cœur demandé. Les dossiers Ant et FOP sont placés en annexes techniques : ils complètent le livrable sans alourdir les chapitres principaux.",
    )
    add_table(
        doc,
        ["Partie", "Dossier source", "Rôle dans le rapport"],
        [
            ("Java RMI", "TP0", "Chapitre principal : objets distants et callbacks."),
            ("XML-RPC", "TP0-XMLRPC", "Chapitre principal : RPC simple sur HTTP/XML."),
            ("SOAP", "TP0-SOAP", "Chapitre principal : enveloppe XML et contrat WSDL."),
            ("REST", "TP0-REST", "Chapitre principal : endpoints HTTP et JSON."),
            ("TCP sockets", "dossier bonus TCP", "Bonus technique : compréhension réseau bas niveau."),
            ("Apache Ant", "ProjetAnt et build.xml", "Annexe : automatisation compilation, documentation et archive."),
            ("Apache FOP", "fop et projet", "Annexe : transformation XML/XSL-FO vers PDF."),
        ],
        weights=[1.4, 1.6, 3.6],
        header_fill=LIGHT_BLUE,
    )


def chapter_rmi(doc, figures):
    chapter(doc, "Chapitre 2 : Application ChatUser avec Java RMI", "Objet distant, registre RMI et callbacks", anchor=ANCHORS["ch2"])
    section_heading(doc, "2.1 Contexte du travail pratique")
    add_paragraph(
        doc,
        "La version RMI du projet montre comment une application Java peut appeler les méthodes d’un objet situé dans une autre JVM. Le serveur publie un objet distant représentant la salle de discussion, tandis que chaque client expose aussi un objet distant permettant de recevoir les messages par callback.",
    )
    section_heading(doc, "2.2 Objectif du travail pratique")
    add_paragraph(
        doc,
        "L’objectif est de réaliser une salle de discussion distribuée en conservant une programmation orientée objet. Le travail attendu consiste à définir les interfaces distantes, implémenter la salle commune, publier l’objet serveur dans le registre RMI et vérifier que plusieurs clients reçoivent bien les messages diffusés.",
    )
    add_paragraph(
        doc,
        "La réponse proposée repose donc sur deux contrats : `ChatRoom` pour les opérations de la salle et `ChatUser` pour le retour vers les clients. Cette double interface permet d’obtenir une communication bidirectionnelle : le client appelle le serveur et le serveur rappelle le client.",
    )
    section_heading(doc, "2.3 Notions utilisées")
    add_paragraph(
        doc,
        "Java RMI repose sur trois idées : une interface distante, une implémentation exportée et un registre qui permet au client de retrouver l’objet distant par son nom. Dans ce projet, `ChatRoom` est l’interface distante côté serveur, tandis que `ChatUser` permet au serveur de rappeler les clients connectés.",
    )
    add_callout(
        doc,
        "Point important",
        "RMI garde une écriture très proche de l’objet local : le client appelle `postMessage`, mais l’exécution réelle se produit côté serveur. Le réseau est masqué par l’infrastructure RMI.",
        "info",
    )
    add_figure(doc, figures, "2.1")
    section_heading(doc, "2.4 Analyse du code")
    add_paragraph(
        doc,
        "L’interface `ChatRoom` définit les opérations accessibles à distance. Elle hérite de `Remote`, et chaque méthode déclare `RemoteException`, ce qui rappelle qu’un appel distant peut échouer pour une raison réseau, une indisponibilité du serveur ou une déconnexion du client.",
    )
    add_code_block(
        doc,
        """
public interface ChatRoom extends Remote {
    void subscribe(ChatUser user, String pseudo) throws RemoteException;
    void unsubscribe(String pseudo) throws RemoteException;
    void postMessage(String pseudo, String message) throws RemoteException;
}
""",
        "Extrait ciblé : interface distante `ChatRoom` dans la version RMI.",
    )
    add_paragraph(
        doc,
        "L’implémentation `ChatRoomImpl` conserve une table de correspondance entre pseudo et objet `ChatUser`. Lorsqu’un message est reçu, la méthode privée `broadcastMessage` parcourt les utilisateurs et appelle `displayMessage` sur chacun. Ce choix correspond à une diffusion active : le serveur pousse le message vers les clients.",
    )
    add_table(
        doc,
        ["Élément", "Rôle", "Justification"],
        [
            ("ChatRoom", "Contrat distant du serveur", "Expose les actions de la salle de discussion."),
            ("ChatUser", "Contrat distant du client", "Autorise le callback `displayMessage`."),
            ("ChatRoomImpl", "État de la salle", "Garde la liste des utilisateurs et diffuse les messages."),
            ("ChatServer", "Publication RMI", "Crée le registre et associe le nom `ChatRoom` à l’objet."),
            ("ChatClient", "Interface utilisateur", "Récupère la référence distante et appelle les méthodes."),
        ],
        weights=[1.4, 2.0, 3.0],
        header_fill=LIGHT_TEAL,
    )
    section_heading(doc, "2.5 Commandes d’exécution")
    add_paragraph(doc, "Le projet possède un fichier `build.xml`. Les commandes Ant évitent de compiler manuellement chaque paquet Java.")
    add_code_block(
        doc,
        """
cd C:\\Users\\DELL\\Downloads\\i-fall\\TP0
ant compile
ant run-server

# dans un autre terminal
ant run-client -Dpseudo=Alice
""",
        "Commandes de compilation et lancement de la version RMI.",
    )
    section_heading(doc, "2.6 Résultat attendu et synthèse")
    add_paragraph(
        doc,
        "Le serveur doit afficher les arrivées, les départs et les messages. Deux clients lancés en parallèle doivent voir les mêmes événements de conversation. La version RMI est très adaptée à un environnement Java homogène, mais elle est moins naturelle si des clients non Java doivent communiquer avec le service.",
    )
    add_paragraph(
        doc,
        "La principale limite est donc l’interopérabilité. En revanche, pour apprendre les objets distribués, RMI est très parlant : la séparation entre interface, implémentation, registre et client montre proprement les responsabilités d’une application distribuée Java.",
    )


def chapter_xmlrpc(doc, figures):
    chapter(doc, "Chapitre 3 : Application ChatUser avec XML-RPC", "Appels de procédures distantes sur HTTP et XML", anchor=ANCHORS["ch3"])
    section_heading(doc, "3.1 Contexte du travail pratique")
    add_paragraph(
        doc,
        "XML-RPC propose une approche différente de RMI. Le client n’obtient pas une référence vers un objet Java distant ; il envoie une requête HTTP contenant le nom d’une méthode et ses paramètres encodés en XML. Le serveur reçoit la requête, identifie la méthode appelée et renvoie une réponse XML.",
    )
    section_heading(doc, "3.2 Objectif du travail pratique")
    add_paragraph(
        doc,
        "L’objectif est de réimplémenter le même besoin ChatUser en remplaçant l’appel objet distant par un appel de procédure distante. Le serveur doit publier des méthodes accessibles par nom, tandis que le client doit construire des appels XML-RPC et interpréter les réponses.",
    )
    add_paragraph(
        doc,
        "La réponse proposée choisit une liste centralisée de messages et un index de lecture. Cette solution évite d’avoir un objet client distant et s’adapte mieux à un protocole requête/réponse : le client revient périodiquement chercher ce qu’il n’a pas encore lu.",
    )
    section_heading(doc, "3.3 Fonctionnement retenu dans ChatUser")
    add_paragraph(
        doc,
        "Dans cette version, le serveur expose un gestionnaire nommé `ChatRoom`. Les opérations principales sont `subscribe`, `unsubscribe`, `postMessage` et `getMessages`. Comme le serveur XML-RPC ne rappelle pas directement les clients, la réception repose sur le polling : le client demande périodiquement les nouveaux messages depuis un index connu.",
    )
    add_figure(doc, figures, "3.1")
    add_callout(
        doc,
        "Différence avec RMI",
        "RMI autorise un callback direct vers l’objet client. XML-RPC reste sur un échange requête/réponse : pour recevoir les messages, le client doit revenir demander ce qui a changé.",
        "warn",
    )
    section_heading(doc, "3.4 Analyse du code")
    add_paragraph(
        doc,
        "Le serveur s’appuie sur `WebServer` de la bibliothèque Apache XML-RPC. Le mapping associe le nom logique `ChatRoom` à la classe `ChatRoomImpl`. Côté métier, la classe conserve une liste de messages et retourne uniquement les messages ajoutés depuis `lastIndex`.",
    )
    add_code_block(
        doc,
        """
WebServer webServer = new WebServer(8080);
XmlRpcServer xmlRpcServer = webServer.getXmlRpcServer();

PropertyHandlerMapping phm = new PropertyHandlerMapping();
phm.addHandler("ChatRoom", ChatRoomImpl.class);
xmlRpcServer.setHandlerMapping(phm);

webServer.start();
""",
        "Extrait ciblé : publication du gestionnaire XML-RPC côté serveur.",
    )
    add_paragraph(
        doc,
        "Le choix de `lastIndex` évite de renvoyer l’historique complet à chaque appel. Le client garde la position du dernier message reçu, puis le serveur calcule la sous-liste à retourner. Ce mécanisme est simple, lisible et suffisant pour un travail pratique de chat.",
    )
    add_table(
        doc,
        ["Opération", "Paramètres", "Réponse attendue", "Rôle"],
        [
            ("subscribe", "pseudo", "OK ou ERROR", "Inscrire un utilisateur."),
            ("postMessage", "pseudo, message", "OK", "Ajouter un message à l’historique."),
            ("getMessages", "lastIndex", "tableau de chaînes", "Retourner les nouveaux messages."),
            ("unsubscribe", "pseudo", "OK", "Retirer l’utilisateur de la salle."),
        ],
        weights=[1.3, 1.7, 1.6, 2.4],
        header_fill=LIGHT_TEAL,
    )
    section_heading(doc, "3.5 Commandes et résultat attendu")
    add_code_block(
        doc,
        """
cd C:\\Users\\DELL\\Downloads\\i-fall\\TP0-XMLRPC
ant compile
ant run-server

# dans un autre terminal
ant run-client
""",
        "Commandes de test de la version XML-RPC.",
    )
    add_paragraph(
        doc,
        "Le serveur doit écouter sur `http://localhost:8080/xmlrpc`. Lorsque deux clients sont lancés, les messages apparaissent après interrogation régulière du serveur. Une légère latence peut exister, car la réception dépend de la fréquence de polling.",
    )
    add_paragraph(
        doc,
        "La limite principale est cette interrogation répétée. Si l’intervalle est trop long, les messages arrivent avec retard ; s’il est trop court, le serveur reçoit beaucoup de requêtes. Le compromis choisi est acceptable pour un travail pratique, car il met en évidence la différence entre callback et polling.",
    )


def chapter_soap(doc, figures):
    chapter(doc, "Chapitre 4 : Application ChatUser avec SOAP", "Enveloppe XML, service HTTP et contrat WSDL", anchor=ANCHORS["ch4"])
    section_heading(doc, "4.1 Contexte du travail pratique")
    add_paragraph(
        doc,
        "SOAP est un protocole de services web basé sur XML. Il impose une structure de message plus formelle que XML-RPC : une enveloppe, un corps et éventuellement un en-tête. Dans le projet, le service est publié sur `/chat` et expose aussi un WSDL simplifié via `/chat?wsdl`.",
    )
    section_heading(doc, "4.2 Objectif du travail pratique")
    add_paragraph(
        doc,
        "L’objectif est de comprendre comment un service web contractuel traite des opérations distantes. Par rapport à XML-RPC, SOAP rend l’échange plus formel : les opérations sont placées dans une enveloppe XML et le service peut fournir une description WSDL.",
    )
    add_paragraph(
        doc,
        "La réponse proposée consiste à garder la logique métier `ChatRoomImpl`, mais à l’envelopper dans une couche SOAP. Ainsi, le cœur de l’application reste comparable aux autres versions, tandis que la couche réseau illustre le traitement XML propre à SOAP.",
    )
    section_heading(doc, "4.3 Architecture du service")
    add_paragraph(
        doc,
        "La version SOAP utilise `HttpServer` pour recevoir des requêtes HTTP. Une requête `GET` avec la query `wsdl` retourne le contrat du service. Une requête `POST` contient une enveloppe SOAP ; le serveur parse le XML avec `SoapUtil`, identifie l’opération demandée, puis appelle la méthode correspondante dans `ChatRoomImpl`.",
    )
    add_figure(doc, figures, "4.1")
    section_heading(doc, "4.4 Analyse du traitement")
    add_paragraph(
        doc,
        "Le code montre bien la séparation entre transport, parsing et logique métier. `ChatServer` reçoit la requête, `SoapUtil` extrait l’opération, puis `ChatRoomImpl` applique la règle fonctionnelle : inscription, désinscription, ajout d’un message ou récupération des messages.",
    )
    add_code_block(
        doc,
        """
if ("GET".equals(exchange.getRequestMethod())
        && "wsdl".equals(exchange.getRequestURI().getQuery())) {
    send(exchange, 200, "text/xml; charset=UTF-8", wsdl());
    return;
}

Document document = SoapUtil.parse(request);
Element operation = SoapUtil.getSoapOperation(document);
""",
        "Extrait ciblé : distinction entre WSDL et traitement d’une opération SOAP.",
    )
    add_paragraph(
        doc,
        "Le WSDL joue le rôle de description du service : il indique le nom du service, l’adresse, les opérations disponibles et le binding SOAP. Même si le WSDL du projet est volontairement compact, il montre l’idée fondamentale : un service SOAP est censé être consommable à partir d’un contrat.",
    )
    add_table(
        doc,
        ["Composant", "Responsabilité", "Observation"],
        [
            ("HttpServer", "Recevoir les requêtes HTTP", "Point d’entrée technique du service."),
            ("SoapUtil", "Construire et parser les enveloppes XML", "Isole la manipulation XML du serveur."),
            ("ChatRoomImpl", "Appliquer la logique de chat", "Reste proche des autres versions."),
            ("WSDL", "Décrire le service", "Aide les clients à connaître les opérations disponibles."),
        ],
        weights=[1.4, 2.4, 2.5],
        header_fill=LIGHT_ORANGE,
    )
    section_heading(doc, "4.5 Commandes et résultat attendu")
    add_code_block(
        doc,
        """
cd C:\\Users\\DELL\\Downloads\\i-fall\\TP0-SOAP
ant compile
ant run-server

# accès au contrat
http://localhost:8081/chat?wsdl
""",
        "Commandes de lancement et vérification du WSDL SOAP.",
    )
    add_paragraph(
        doc,
        "La page WSDL doit être accessible dans un navigateur. Les clients doivent ensuite pouvoir envoyer des messages et récupérer les nouveaux messages à travers des enveloppes SOAP. Cette version est plus verbeuse que REST, mais plus contractuelle.",
    )
    add_paragraph(
        doc,
        "La limite de SOAP dans ce projet est la lourdeur des messages XML. Pour un petit chat, REST est plus léger. Mais SOAP reste intéressant pédagogiquement, car il montre la logique des services fortement structurés, encore présents dans certains systèmes d’information.",
    )


def chapter_rest(doc, figures):
    chapter(doc, "Chapitre 5 : Application ChatUser avec REST", "Endpoints HTTP, JSON et ressources applicatives", anchor=ANCHORS["ch5"])
    section_heading(doc, "5.1 Contexte du travail pratique")
    add_paragraph(
        doc,
        "La version REST traduit le fonctionnement de ChatUser en endpoints HTTP. Les opérations ne sont plus exposées comme des objets distants ou comme des procédures XML, mais comme des points d’accès web : inscription, envoi de message, consultation des nouveaux messages et désinscription.",
    )
    section_heading(doc, "5.2 Objectif du travail pratique")
    add_paragraph(
        doc,
        "L’objectif est d’exprimer l’application sous forme d’API web simple. Le client n’a plus besoin d’une bibliothèque RPC particulière : il envoie des requêtes HTTP et reçoit du JSON. Ce modèle facilite les tests avec des outils standards et prépare à la conception d’API modernes.",
    )
    add_paragraph(
        doc,
        "La réponse proposée définit quatre endpoints, chacun associé à une responsabilité précise. Le serveur vérifie la méthode HTTP, lit les données envoyées et répond avec un statut JSON. Le comportement métier reste celui du chat, mais l’interface devient web.",
    )
    section_heading(doc, "5.3 Endpoints disponibles")
    add_paragraph(
        doc,
        "Le serveur REST écoute sur le port 8082. Les échanges sont en JSON, ce qui rend les requêtes plus légères et lisibles que les enveloppes SOAP. Le client Swing peut utiliser ces endpoints, mais ils peuvent aussi être testés avec un navigateur, curl ou Postman.",
    )
    add_figure(doc, figures, "5.1")
    add_table(
        doc,
        ["Endpoint", "Méthode", "Corps / paramètres", "Rôle"],
        [
            ("/chat/subscribe", "POST", "{\"pseudo\":\"Alice\"}", "Inscrire un utilisateur."),
            ("/chat/message", "POST", "{\"pseudo\":\"Alice\",\"message\":\"Bonjour\"}", "Publier un message."),
            ("/chat/messages", "GET", "lastIndex=0", "Lire les nouveaux messages."),
            ("/chat/unsubscribe", "POST", "{\"pseudo\":\"Alice\"}", "Quitter la salle."),
        ],
        weights=[1.8, 0.9, 2.5, 2.2],
        header_fill=LIGHT_TEAL,
    )
    section_heading(doc, "5.4 Analyse du code")
    add_paragraph(
        doc,
        "La classe `ChatServer` crée un contexte HTTP pour chaque endpoint. Les méthodes `handleSubscribe`, `handleMessage` ou `handleMessages` vérifient la méthode HTTP, lisent le JSON, appellent la logique métier puis renvoient une réponse JSON avec un statut explicite.",
    )
    add_code_block(
        doc,
        """
server.createContext("/chat/subscribe", ChatServer::handleSubscribe);
server.createContext("/chat/message", ChatServer::handleMessage);
server.createContext("/chat/messages", ChatServer::handleMessages);

exchange.getResponseHeaders()
        .set("Content-Type", "application/json; charset=UTF-8");
""",
        "Extrait ciblé : publication des endpoints REST et réponse JSON.",
    )
    add_callout(
        doc,
        "Choix pédagogique",
        "Le projet utilise un `JsonUtil` simple au lieu d’une grosse bibliothèque externe. Cela rend le mécanisme visible : lecture du corps, extraction des champs et construction manuelle de la réponse.",
        "info",
    )
    section_heading(doc, "5.5 Commandes et résultat attendu")
    add_code_block(
        doc,
        """
cd C:\\Users\\DELL\\Downloads\\i-fall\\TP0-REST
ant compile
ant run-server

# exemple de test conceptuel
POST http://localhost:8082/chat/subscribe
{"pseudo":"Alice"}
""",
        "Commandes et exemple de requête pour la version REST.",
    )
    add_paragraph(
        doc,
        "Le serveur doit afficher la liste des endpoints disponibles. Les clients doivent pouvoir s’inscrire, publier des messages et récupérer les messages avec `lastIndex`. REST est la version la plus proche des API web modernes.",
    )
    add_paragraph(
        doc,
        "La limite de cette version est que le projet reste volontairement simple : il n’y a pas d’authentification, pas de persistance et pas de vraie gestion d’erreurs avancée. Pour un travail pratique, cette simplicité est utile car elle laisse apparaître clairement la relation entre URL, méthode HTTP, JSON et logique métier.",
    )


def chapter_comparison(doc, figures):
    chapter(doc, "Chapitre 6 : Comparaison générale des architectures", "Lire les différences techniques et pédagogiques", anchor=ANCHORS["ch6"])
    section_heading(doc, "6.1 Vue comparative")
    add_paragraph(
        doc,
        "Les quatre versions principales résolvent le même problème, mais elles ne le modélisent pas de la même manière. RMI privilégie l’objet distant, XML-RPC privilégie l’appel de procédure simple, SOAP formalise les échanges XML avec un contrat, et REST transforme les actions en ressources HTTP manipulées avec JSON.",
    )
    add_figure(doc, figures, "6.1")
    section_heading(doc, "6.2 Tableau de synthèse")
    add_table(
        doc,
        ["Critère", "RMI", "XML-RPC", "SOAP", "REST"],
        [
            ("Modèle", "Objet distant", "Procédure distante", "Service contractuel", "Ressource HTTP"),
            ("Format", "Sérialisation Java", "XML", "XML SOAP", "JSON"),
            ("Transport", "RMI/JRMP", "HTTP", "HTTP", "HTTP"),
            ("Réception messages", "Callback serveur", "Polling", "Polling", "Polling"),
            ("Interopérabilité", "Faible hors Java", "Bonne", "Bonne mais verbeuse", "Très bonne"),
            ("Lisibilité réseau", "Peu lisible", "Lisible mais XML", "Très verbeux", "Lisible et compact"),
            ("Usage typique", "Système Java homogène", "RPC léger", "SI formel/contrat", "API web moderne"),
        ],
        weights=[1.4, 1.4, 1.5, 1.5, 1.6],
        header_fill=LIGHT_BLUE,
        font_size=8.2,
    )
    section_heading(doc, "6.3 Analyse")
    add_paragraph(
        doc,
        "La différence la plus structurante concerne le sens de communication. RMI permet au serveur d’appeler directement les clients grâce à `ChatUser`. Les autres versions principales ne disposent pas de ce callback ; le client doit donc interroger le serveur pour récupérer les nouveaux messages. Cette distinction explique une grande partie des différences de code.",
    )
    add_paragraph(
        doc,
        "Le deuxième axe concerne le format. SOAP et XML-RPC utilisent XML, ce qui facilite la représentation structurée mais augmente la verbosité. REST utilise JSON, plus compact et plus habituel dans les API modernes. RMI, lui, masque le format réseau au développeur Java.",
    )
    add_callout(
        doc,
        "Synthèse pédagogique",
        "RMI montre l’appel objet distant, XML-RPC montre le RPC simple, SOAP montre le service web contractuel, REST montre l’API web moderne. Le projet donne donc une comparaison complète des styles de communication distribuée.",
        "ok",
    )


def chapter_validation(doc, figures):
    chapter(doc, "Chapitre 7 : Guide d’exécution et de validation", "Compiler, lancer, capturer et contrôler les résultats", anchor=ANCHORS["ch7"])
    section_heading(doc, "7.1 Préparation de l’environnement")
    add_paragraph(
        doc,
        "Les projets sont des applications Java organisées avec Apache Ant. Avant de tester, il faut vérifier que Java et Ant sont disponibles dans le terminal. Ensuite, chaque version se lance selon le même principe : compilation, démarrage du serveur, lancement d’un ou plusieurs clients.",
    )
    add_code_block(
        doc,
        """
java -version
ant -version
""",
        "Contrôle minimal de l’environnement Java et Ant.",
    )
    add_figure(doc, figures, "7.1")
    section_heading(doc, "7.2 Scénario de test commun")
    add_paragraph(
        doc,
        "Le scénario de validation doit rester identique pour toutes les versions principales. On lance un serveur, puis deux clients. Le premier client utilise le pseudo Alice, le second utilise Bob. Alice envoie un message, Bob doit le recevoir ; Bob répond, Alice doit le recevoir à son tour.",
    )
    add_bullets(
        doc,
        [
            "Le serveur démarre sans erreur sur le port prévu.",
            "Alice peut rejoindre la salle de discussion.",
            "Bob peut rejoindre la même salle.",
            "Les notifications d’arrivée sont visibles.",
            "Un message envoyé par Alice apparaît chez Bob.",
            "Un message envoyé par Bob apparaît chez Alice.",
            "La sortie d’un client ne bloque pas le serveur.",
        ],
    )
    section_heading(doc, "7.3 Commandes par version")
    add_table(
        doc,
        ["Version", "Dossier", "Port", "Commandes principales"],
        [
            ("RMI", "TP0", "1099", "ant compile ; ant run-server ; ant run-client -Dpseudo=Alice"),
            ("XML-RPC", "TP0-XMLRPC", "8080", "ant compile ; ant run-server ; ant run-client"),
            ("SOAP", "TP0-SOAP", "8081", "ant compile ; ant run-server ; ant run-client"),
            ("REST", "TP0-REST", "8082", "ant compile ; ant run-server ; ant run-client"),
        ],
        weights=[1.0, 1.4, 0.8, 4.0],
        header_fill=LIGHT_TEAL,
    )
    section_heading(doc, "7.4 Captures d’exécution à insérer")
    add_paragraph(
        doc,
        "Les captures ne doivent pas être placées en rafale. Elles servent à prouver l’exécution et doivent apparaître après une explication du scénario. Les cadres ci-dessous indiquent les captures recommandées lorsque l’application est lancée localement.",
    )
    add_capture_placeholder(
        doc,
        "Capture 1 : serveur ChatUser démarré",
        "Insérer une capture du terminal serveur après `ant run-server`, avec le port visible et le message indiquant que le serveur attend les connexions.",
    )
    add_capture_placeholder(
        doc,
        "Capture 2 : deux clients connectés",
        "Insérer une capture montrant deux fenêtres client avec des pseudos différents, par exemple Alice et Bob.",
    )
    add_capture_placeholder(
        doc,
        "Capture 3 : message envoyé et reçu",
        "Insérer une capture où un message envoyé par un client apparaît chez l’autre, pour valider la diffusion ou la récupération par polling.",
    )
    add_capture_placeholder(
        doc,
        "Capture 4 : sortie d’un utilisateur",
        "Insérer une capture montrant la notification de départ ou l’absence d’erreur serveur après fermeture d’un client.",
    )
    section_heading(doc, "7.5 Points de contrôle")
    add_paragraph(
        doc,
        "La validation ne doit pas seulement consister à voir que le programme se lance. Il faut vérifier le comportement attendu : unicité des pseudos, réception des messages, absence de blocage serveur, gestion correcte des clients déconnectés et cohérence entre les messages affichés par les différents clients.",
    )
    add_callout(
        doc,
        "Bon réflexe de test",
        "Toujours garder le terminal serveur visible pendant les tests : il indique les connexions, les messages reçus et les erreurs éventuelles.",
        "method",
    )


def bonus_tcp(doc, figures):
    chapter(doc, "Bonus technique : variante ChatUser avec sockets TCP", "Compréhension réseau bas niveau", anchor=ANCHORS["bonus_tcp"])
    section_heading(doc, "B.1 Pourquoi cette partie est en bonus")
    add_paragraph(
        doc,
        "La version TCP n’est pas intégrée aux chapitres principaux parce qu’elle ne faisait pas partie du périmètre demandé. Elle reste cependant intéressante : elle expose directement le mécanisme réseau utilisé à bas niveau par beaucoup de technologies plus abstraites.",
    )
    add_figure(doc, figures, "B.1")
    section_heading(doc, "B.2 Fonctionnement")
    add_paragraph(
        doc,
        "Le serveur ouvre un `ServerSocket` sur le port 8083. À chaque connexion entrante, il crée un `ClientHandler` exécuté dans un thread séparé. Le client envoie d’abord son pseudo, puis chaque ligne suivante est interprétée comme un message de chat.",
    )
    add_table(
        doc,
        ["Étape", "Message ou action", "Traitement côté serveur"],
        [
            ("1", "Ouverture d’une socket vers localhost:8083", "Le `ServerSocket` accepte la connexion entrante."),
            ("2", "Envoi du pseudo sur la première ligne", "Le serveur valide le pseudo et répond `OK` ou `ERROR`."),
            ("3", "Envoi d’une ligne de texte", "La ligne est considérée comme un message de chat."),
            ("4", "Diffusion", "`ChatRoomImpl` parcourt les clients connectés et appelle `sendMessage`."),
            ("5", "Déconnexion", "Le client est retiré de la liste et les autres utilisateurs sont notifiés."),
        ],
        weights=[0.8, 2.6, 3.6],
        header_fill=LIGHT_TEAL,
    )
    add_code_block(
        doc,
        """
try (ServerSocket serverSocket = new ServerSocket(8083)) {
    while (true) {
        Socket socket = serverSocket.accept();
        ClientHandler handler = new ClientHandler(socket);
        new Thread(handler).start();
    }
}
""",
        "Extrait ciblé : acceptation des connexions TCP et création d’un thread par client.",
    )
    add_paragraph(
        doc,
        "Cette approche est plus manuelle : il faut gérer le protocole ligne par ligne, les connexions, les déconnexions et la diffusion. Elle est utile pour apprendre, mais elle demande plus de discipline qu’un service HTTP ou qu’un appel distant déjà structuré.",
    )
    section_heading(doc, "B.3 Analyse et limites")
    add_paragraph(
        doc,
        "La variante TCP rend visibles les détails que les autres technologies cachent : ouverture de socket, flux d’entrée/sortie, boucle d’acceptation et concurrence. Cette transparence est intéressante pour apprendre le réseau, mais elle impose de définir soi-même un protocole applicatif. Ici, le protocole est volontairement simple : une ligne pour le pseudo, puis une ligne par message.",
    )
    add_paragraph(
        doc,
        "Les limites apparaissent vite dans une application réelle : pas de format structuré comme JSON ou XML, pas de contrat de service, pas de gestion avancée des erreurs et un risque de complexité dès qu’il faut ajouter des commandes plus riches. Pour cette raison, TCP est conservé comme bonus pédagogique plutôt que comme solution principale du rapport.",
    )


def annexes(doc, figures):
    chapter(doc, "Annexes techniques", "Automatisation Ant, génération PDF FOP et références", anchor=ANCHORS["annexes"])
    section_heading(doc, "Annexe A.1 : Projet Ant")
    add_paragraph(
        doc,
        "Apache Ant automatise les tâches répétitives d’un projet Java : création des répertoires, compilation, génération de documentation, création d’une archive JAR et nettoyage. Le dossier `ProjetAnt` illustre cette logique avec un programme simple `Bonjour.java`.",
    )
    add_figure(doc, figures, "A.1")
    add_table(
        doc,
        ["Cible Ant", "Rôle", "Résultat attendu"],
        [
            ("init", "Créer les répertoires nécessaires", "`bin`, `doc` et `archive` sont disponibles."),
            ("compile", "Compiler les sources Java", "Les fichiers `.class` sont générés dans `bin`."),
            ("doc", "Produire la documentation Javadoc", "La documentation HTML est générée dans `doc`."),
            ("arch", "Créer une archive exécutable", "Un fichier JAR est produit dans `archive`."),
            ("clean", "Nettoyer les sorties générées", "Le dossier `bin` est supprimé pour repartir proprement."),
            ("run", "Exécuter le programme", "La classe `Bonjour` est lancée via Ant."),
            ("all", "Enchaîner le build complet", "Nettoyage, compilation, documentation et archive sont réalisés."),
        ],
        weights=[1.1, 2.7, 3.2],
        header_fill=LIGHT_BLUE,
    )
    add_code_block(
        doc,
        """
<target name="all" depends="clean,init,compile,doc,arch">
    <echo message="Build complet terminé avec succès !"/>
</target>
""",
        "Extrait ciblé : cible Ant globale du projet.",
    )
    add_paragraph(
        doc,
        "L’intérêt d’Ant est la reproductibilité : au lieu d’expliquer une longue série de commandes, le projet documente les actions dans `build.xml`. Un autre étudiant peut relancer les mêmes étapes avec les mêmes noms de cibles.",
    )
    add_paragraph(
        doc,
        "Dans le cadre du projet i-fall, cette annexe explique aussi pourquoi chaque dossier ChatUser contient un `build.xml`. Le rapport met donc en relation les commandes utilisées dans les chapitres et le mécanisme d’automatisation qui les rend reproductibles.",
    )
    section_heading(doc, "Annexe A.2 : Génération PDF avec Apache FOP")
    add_paragraph(
        doc,
        "Apache FOP transforme des documents XSL-FO en sorties imprimables comme PDF. Dans le dossier `fop`, le fichier `equipe.xml` contient les données, tandis que `statistiques-fo.xsl` décrit la mise en forme. Le dossier `projet` contient aussi un exemple Java `Fo2Pdf.java` qui pilote FOP côté code.",
    )
    add_figure(doc, figures, "A.2")
    add_table(
        doc,
        ["Fichier / dossier", "Rôle", "Explication"],
        [
            ("equipe.xml", "Source de données", "Contient les équipes, victoires, défaites et classements."),
            ("statistiques.xsl", "Transformation HTML/XML", "Présente une transformation lisible côté navigateur ou XML."),
            ("statistiques-fo.xsl", "Transformation XSL-FO", "Décrit la mise en page destinée au rendu PDF."),
            ("test.fo", "Document XSL-FO direct", "Sert de fichier d’entrée simple pour tester FOP."),
            ("Fo2Pdf.java", "Programme Java", "Configure le driver FOP et produit le PDF à partir du fichier FO."),
            ("out/test.pdf", "Sortie générée", "Montre le résultat final après conversion."),
        ],
        weights=[1.6, 1.8, 3.6],
        header_fill=LIGHT_TEAL,
    )
    add_code_block(
        doc,
        """
Driver driver = new Driver();
driver.setRenderer(Driver.RENDER_PDF);
driver.setInputSource(new InputSource(in));
driver.setOutputStream(out);
driver.run();
""",
        "Extrait ciblé : conversion XSL-FO vers PDF avec Apache FOP.",
    )
    add_paragraph(
        doc,
        "La séparation entre XML et XSL-FO est importante : les données restent propres et indépendantes, tandis que la feuille de transformation décide de la présentation finale. C’est une logique proche des architectures qui séparent modèle, traitement et rendu.",
    )
    add_paragraph(
        doc,
        "Cette annexe complète les travaux sur les services distribués en montrant une autre forme de chaîne technique : une donnée structurée est transformée en document final. Le principe général reste similaire à une architecture propre : séparer l’information, la transformation et le format de sortie.",
    )
    section_heading(doc, "Annexe A.3 : Références consultées")
    add_table(
        doc,
        ["Référence", "Lien", "Utilisation dans le rapport"],
        [
            ("Oracle Java RMI API Guide", "https://docs.oracle.com/en/java/javase/11/rmi/index.html", "Définition et logique de Java RMI."),
            ("Apache XML-RPC", "https://ws.apache.org/xmlrpc/", "Description du protocole XML-RPC et de l’implémentation Java."),
            ("W3C SOAP 1.1", "https://www.w3.org/TR/SOAP/", "Structure des messages SOAP."),
            ("W3C WSDL", "https://www.w3.org/TR/wsdl.html", "Rôle du contrat WSDL."),
            ("MDN REST", "https://developer.mozilla.org/en-US/docs/Glossary/REST", "Rappel sur REST comme style architectural."),
            ("Apache Ant Manual", "https://ant.apache.org/manual/", "Cibles et automatisation de build."),
            ("Apache FOP Quick Start", "https://xmlgraphics.apache.org/fop/quickstartguide.html", "Principe de génération PDF avec FOP."),
        ],
        weights=[1.8, 2.9, 2.3],
        header_fill=LIGHT_BLUE,
        font_size=7.8,
    )
    section_heading(doc, "Conclusion générale")
    add_paragraph(
        doc,
        "Le projet i-fall permet de comprendre que la communication distribuée ne dépend pas seulement du code métier, mais aussi du modèle d’échange choisi. RMI donne une vision objet, XML-RPC simplifie l’appel distant, SOAP formalise les contrats XML et REST propose une approche web lisible et interopérable. Les annexes Ant et FOP complètent le travail en montrant l’automatisation de build et la transformation documentaire.",
    )
    add_callout(
        doc,
        "Bilan",
        "Le livrable final met donc en relation le code, les architectures, les commandes de test et les choix techniques, afin que le lecteur comprenne le projet au lieu de seulement voir une liste de fichiers.",
        "ok",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspace(p, before=8, after=0)
    r = p.add_run("Code source du projet : ")
    set_font(r, size=9.2, bold=True, color=MUTED)
    add_external_hyperlink(p, REPO_URL, REPO_URL, size=9.2)


def build_doc():
    figures = build_figures()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    chapter_intro(doc, figures)
    chapter_rmi(doc, figures)
    chapter_xmlrpc(doc, figures)
    chapter_soap(doc, figures)
    chapter_rest(doc, figures)
    chapter_comparison(doc, figures)
    chapter_validation(doc, figures)
    bonus_tcp(doc, figures)
    annexes(doc, figures)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
